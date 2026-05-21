# backend/main.py

from fastapi import FastAPI, UploadFile, File, HTTPException, Depends, status, Form
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from sqlalchemy import create_engine, text, func
from sqlalchemy.orm import sessionmaker, Session
from sentence_transformers import SentenceTransformer
from datetime import datetime, timedelta
from typing import List, Optional, Dict, Any
import pypdf
from docx import Document as DocxDocument
import ollama
import google.generativeai as genai
import os
import shutil
import json
import uuid
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
from pydantic import BaseModel, EmailStr
from jose import JWTError, jwt
from passlib.context import CryptContext
from dotenv import load_dotenv
import stripe
import razorpay
from twilio.rest import Client as TwilioClient
import numpy as np

# Import models
from models import *

# ============ CONFIGURATION ============
load_dotenv()

# Database
DATABASE_URL = os.getenv("DATABASE_URL")
if not DATABASE_URL:
    raise ValueError("No DATABASE_URL set in .env file")

# AI Configuration
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "gemma2:2b-instruct-q8_0")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "nomic-ai/nomic-embed-text-v1")
EMBEDDING_DIMENSION = int(os.getenv("EMBEDDING_DIMENSION", 768))

# RAG Configuration
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", 800))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", 100))
TOP_K_RESULTS = int(os.getenv("TOP_K_RESULTS", 8))

# File Storage
UPLOAD_DIR = os.getenv("UPLOAD_DIR", "uploads")
ICONS_DIR = os.getenv("ICONS_DIR", "icons")
DOCUMENTS_DIR = os.getenv("DOCUMENTS_DIR", "documents")

# Security
SECRET_KEY = os.getenv("SECRET_KEY", "fallback-secret-key-DO-NOT-USE-IN-PROD")
ALGORITHM = os.getenv("ALGORITHM", "HS256")
ACCESS_TOKEN_EXPIRE_MINUTES = int(os.getenv("ACCESS_TOKEN_EXPIRE_MINUTES", 30))

# Payment Configuration
STRIPE_SECRET_KEY = os.getenv("STRIPE_SECRET_KEY")
RAZORPAY_KEY_ID = os.getenv("RAZORPAY_KEY_ID")
RAZORPAY_KEY_SECRET = os.getenv("RAZORPAY_KEY_SECRET")

# Communication
TWILIO_ACCOUNT_SID = os.getenv("TWILIO_ACCOUNT_SID")
TWILIO_AUTH_TOKEN = os.getenv("TWILIO_AUTH_TOKEN")
TWILIO_PHONE_NUMBER = os.getenv("TWILIO_PHONE_NUMBER")

# Initialize services
PWD_CONTEXT = CryptContext(schemes=["argon2"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")

# Initialize payment processors
if STRIPE_SECRET_KEY:
    stripe.api_key = STRIPE_SECRET_KEY

if RAZORPAY_KEY_ID and RAZORPAY_KEY_SECRET:
    razorpay_client = razorpay.Client(auth=(RAZORPAY_KEY_ID, RAZORPAY_KEY_SECRET))

# Initialize Twilio
if TWILIO_ACCOUNT_SID and TWILIO_AUTH_TOKEN:
    twilio_client = TwilioClient(TWILIO_ACCOUNT_SID, TWILIO_AUTH_TOKEN)

# Initialize Gemini
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)

# ============ DATABASE SETUP ============
engine = create_engine(DATABASE_URL)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

# Create directories
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(ICONS_DIR, exist_ok=True)
os.makedirs(DOCUMENTS_DIR, exist_ok=True)

# Initialize embedding model
try:
    embedding_model = SentenceTransformer(EMBEDDING_MODEL)
except Exception as e:
    print(f"Warning: Could not load embedding model: {e}")
    embedding_model = None

# ============ PYDANTIC MODELS ============

class UserCreate(BaseModel):
    username: str
    email: EmailStr
    password: str
    full_name: str
    phone: Optional[str] = None
    role: UserRole = UserRole.CLIENT

class UserResponse(BaseModel):
    id: int
    username: str
    email: str
    full_name: str
    role: UserRole
    status: UserStatus
    is_verified: bool

class Token(BaseModel):
    access_token: str
    token_type: str

class QueryRequest(BaseModel):
    question: str
    document_id: Optional[int] = None
    case_id: Optional[int] = None
    chat_history: List[Dict[str, str]] = []
    session_id: Optional[str] = None
    visitor_info: Optional[Dict[str, Any]] = None

class QueryResponse(BaseModel):
    answer: str
    sources: List[str]
    agent_used: Optional[str] = None
    session_id: str

class CaseCreate(BaseModel):
    title: str
    description: str
    case_type: str
    client_id: int
    priority: CasePriority = CasePriority.MEDIUM
    deadline: Optional[datetime] = None

class AppointmentCreate(BaseModel):
    title: str
    description: Optional[str] = None
    client_id: int
    case_id: Optional[int] = None
    scheduled_at: datetime
    duration_minutes: int = 60
    meeting_type: str = "video"

class PaymentCreate(BaseModel):
    amount: float
    currency: str = "USD"
    description: str
    payment_method: str = "stripe"
    case_id: Optional[int] = None
    appointment_id: Optional[int] = None

class AgentCreate(BaseModel):
    name: str
    alias: str
    model: str
    system_prompt: str
    temperature: float = 0.7
    specialization: AgentSpecialization = AgentSpecialization.GENERAL

class FilterCreate(BaseModel):
    keyword: str
    filter_type: FilterType
    action: Optional[str] = None

# ============ FASTAPI APP ============
app = FastAPI(
    title="Unified Legal Technology Platform",
    description="Enterprise-grade legal RAG chatbot with case management",
    version="1.0.0"
)

# CORS Configuration
allowed_origins = os.getenv("ALLOWED_ORIGINS", "http://localhost:3000,http://localhost:3001").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ============ DEPENDENCY FUNCTIONS ============

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

def verify_password(plain_password: str, hashed_password: str) -> bool:
    return PWD_CONTEXT.verify(plain_password, hashed_password)

def get_password_hash(password: str) -> str:
    return PWD_CONTEXT.hash(password)

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

async def get_current_user(token: str = Depends(oauth2_scheme), db: Session = Depends(get_db)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise credentials_exception
    except JWTError:
        raise credentials_exception
    
    user = db.query(User).filter(User.username == username).first()
    if user is None:
        raise credentials_exception
    return user

def get_embeddings(text: str) -> List[float]:
    """Generate embeddings using available model"""
    if embedding_model:
        return embedding_model.encode(text).tolist()
    elif GEMINI_API_KEY:
        try:
            result = genai.embed_content(model="models/embedding-001", content=text)
            return result["embedding"]
        except Exception as e:
            print(f"Gemini embedding error: {e}")
            return [0.0] * EMBEDDING_DIMENSION
    else:
        return [0.0] * EMBEDDING_DIMENSION

def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """Split text into overlapping chunks"""
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        chunks.append(chunk)
        start = end - overlap
        if start >= len(text):
            break
    return chunks

def extract_text_from_file(file_path: str, filename: str) -> str:
    """Extract text from various file formats"""
    try:
        if filename.lower().endswith('.pdf'):
            reader = pypdf.PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        elif filename.lower().endswith('.docx'):
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        elif filename.lower().endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            raise ValueError(f"Unsupported file format: {filename}")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting text: {str(e)}")

def get_llm_response(prompt: str, model: str = OLLAMA_MODEL, use_gemini: bool = False) -> str:
    """Get response from LLM (Ollama or Gemini)"""
    try:
        if use_gemini and GEMINI_API_KEY:
            model_instance = genai.GenerativeModel('gemini-1.5-flash')
            response = model_instance.generate_content(prompt)
            return response.text
        else:
            response = ollama.chat(
                model=model,
                messages=[{"role": "user", "content": prompt}]
            )
            return response['message']['content']
    except Exception as e:
        return f"Error generating response: {str(e)}"

def check_message_filters(message: str, db: Session) -> Optional[Dict[str, Any]]:
    """Check if message triggers any filters"""
    filters = db.query(MessageFilter).filter(MessageFilter.is_active == True).all()
    
    for filter_obj in filters:
        if filter_obj.keyword.lower() in message.lower():
            return {
                "filter_type": filter_obj.filter_type.value,
                "action": filter_obj.action,
                "keyword": filter_obj.keyword
            }
    return None

def select_agent(message: str, db: Session) -> Optional[AIAgent]:
    """Select appropriate AI agent based on message content"""
    agents = db.query(AIAgent).filter(AIAgent.is_active == True).all()
    
    # Simple keyword-based agent selection
    legal_keywords = {
        AgentSpecialization.CONTRACT_ANALYSIS: ["contract", "agreement", "terms", "clause"],
        AgentSpecialization.LITIGATION: ["lawsuit", "court", "trial", "litigation"],
        AgentSpecialization.CORPORATE: ["corporate", "business", "company", "merger"],
        AgentSpecialization.FAMILY_LAW: ["divorce", "custody", "family", "marriage"],
        AgentSpecialization.CRIMINAL_LAW: ["criminal", "crime", "arrest", "charges"],
        AgentSpecialization.IMMIGRATION: ["immigration", "visa", "citizenship", "green card"],
        AgentSpecialization.INTELLECTUAL_PROPERTY: ["patent", "trademark", "copyright", "IP"],
        AgentSpecialization.REAL_ESTATE: ["property", "real estate", "lease", "mortgage"]
    }
    
    message_lower = message.lower()
    for specialization, keywords in legal_keywords.items():
        if any(keyword in message_lower for keyword in keywords):
            agent = next((a for a in agents if a.specialization == specialization), None)
            if agent:
                return agent
    
    # Default to general agent
    return next((a for a in agents if a.specialization == AgentSpecialization.GENERAL), None)

# ============ AUTHENTICATION ENDPOINTS ============

@app.post("/register", response_model=UserResponse)
async def register_user(user_data: UserCreate, db: Session = Depends(get_db)):
    # Check if user exists
    if db.query(User).filter(User.username == user_data.username).first():
        raise HTTPException(status_code=400, detail="Username already registered")
    if db.query(User).filter(User.email == user_data.email).first():
        raise HTTPException(status_code=400, detail="Email already registered")
    
    # Create user
    hashed_password = get_password_hash(user_data.password)
    db_user = User(
        username=user_data.username,
        email=user_data.email,
        hashed_password=hashed_password,
        full_name=user_data.full_name,
        phone=user_data.phone,
        role=user_data.role
    )
    db.add(db_user)
    db.commit()
    db.refresh(db_user)
    
    return UserResponse(
        id=db_user.id,
        username=db_user.username,
        email=db_user.email,
        full_name=db_user.full_name,
        role=db_user.role,
        status=db_user.status,
        is_verified=db_user.is_verified
    )

@app.post("/token", response_model=Token)
async def login_for_access_token(form_data: OAuth2PasswordRequestForm = Depends(), db: Session = Depends(get_db)):
    user = db.query(User).filter(User.username == form_data.username).first()
    if not user or not verify_password(form_data.password, user.hashed_password):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user.username}, expires_delta=access_token_expires
    )
    return {"access_token": access_token, "token_type": "bearer"}

@app.get("/users/me", response_model=UserResponse)
async def read_users_me(current_user: User = Depends(get_current_user)):
    return UserResponse(
        id=current_user.id,
        username=current_user.username,
        email=current_user.email,
        full_name=current_user.full_name,
        role=current_user.role,
        status=current_user.status,
        is_verified=current_user.is_verified
    )

# ============ DOCUMENT MANAGEMENT ============

@app.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    case_id: Optional[int] = Form(None),
    is_confidential: bool = Form(False),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # Validate file type
    allowed_extensions = {'.pdf', '.docx', '.txt'}
    file_extension = os.path.splitext(file.filename)[1].lower()
    if file_extension not in allowed_extensions:
        raise HTTPException(status_code=400, detail="Unsupported file type")
    
    # Generate unique filename
    unique_filename = f"{uuid.uuid4()}{file_extension}"
    file_path = os.path.join(DOCUMENTS_DIR, unique_filename)
    
    # Save file
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    # Extract text
    text_content = extract_text_from_file(file_path, file.filename)
    
    # Create document record
    document = Document(
        filename=unique_filename,
        original_filename=file.filename,
        file_path=file_path,
        file_size=os.path.getsize(file_path),
        mime_type=file.content_type,
        uploader_id=current_user.id,
        case_id=case_id,
        is_confidential=is_confidential
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    
    # Process text into chunks and embeddings
    chunks = chunk_text(text_content)
    for i, chunk in enumerate(chunks):
        embedding = get_embeddings(chunk)
        chunk_record = DocumentChunk(
            document_id=document.id,
            chunk_text=chunk,
            chunk_index=i,
            embedding=embedding
        )
        db.add(chunk_record)
    
    db.commit()
    
    return {
        "id": document.id,
        "filename": document.original_filename,
        "upload_date": document.upload_date,
        "chunk_count": len(chunks)
    }

@app.get("/documents")
async def list_documents(
    case_id: Optional[int] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    query = db.query(Document)
    
    # Filter by case if specified
    if case_id:
        query = query.filter(Document.case_id == case_id)
    
    # Filter by access permissions
    if current_user.role not in [UserRole.SUPER_ADMIN, UserRole.ADMIN]:
        query = query.filter(Document.uploader_id == current_user.id)
    
    documents = query.all()
    
    return [
        {
            "id": doc.id,
            "filename": doc.original_filename,
            "upload_date": doc.upload_date,
            "case_id": doc.case_id,
            "is_confidential": doc.is_confidential,
            "chunk_count": len(doc.chunks)
        }
        for doc in documents
    ]

# ============ RAG QUERY ENDPOINT ============

@app.post("/query", response_model=QueryResponse)
async def process_query(
    query_request: QueryRequest,
    db: Session = Depends(get_db)
):
    # Generate session ID if not provided
    session_id = query_request.session_id or str(uuid.uuid4())
    
    # Create or update chat session
    session = db.query(ChatSession).filter(ChatSession.session_id == session_id).first()
    if not session:
        visitor_info = query_request.visitor_info or {}
        session = ChatSession(
            session_id=session_id,
            visitor_name=visitor_info.get("name"),
            visitor_email=visitor_info.get("email"),
            visitor_ip=visitor_info.get("ip"),
            visitor_location=visitor_info.get("location"),
            visitor_device=visitor_info.get("device"),
            visitor_browser=visitor_info.get("browser"),
            case_id=query_request.case_id
        )
        db.add(session)
    
    # Check message filters
    filter_result = check_message_filters(query_request.question, db)
    if filter_result and filter_result["filter_type"] == "block":
        # Log blocked message
        message = ChatMessage(
            session_id=session_id,
            role="user",
            message=query_request.question,
            filtered=f"Blocked: {filter_result['keyword']}"
        )
        db.add(message)
        db.commit()
        
        return QueryResponse(
            answer="I'm sorry, but I cannot process that request due to content policy.",
            sources=[],
            session_id=session_id
        )
    
    # Select appropriate agent
    agent = select_agent(query_request.question, db)
    agent_name = agent.name if agent else "default"
    
    # Generate query embedding
    query_embedding = get_embeddings(query_request.question)
    
    # Search for relevant document chunks
    search_query = text("""
        SELECT dc.chunk_text, d.original_filename
        FROM document_chunks dc
        JOIN documents d ON dc.document_id = d.id
        WHERE (:case_id IS NULL OR d.case_id = :case_id)
        ORDER BY dc.embedding <-> :query_embedding
        LIMIT :top_k
    """)
    
    results = db.execute(search_query, {
        "query_embedding": str(query_embedding),
        "case_id": query_request.case_id,
        "top_k": TOP_K_RESULTS
    }).fetchall()
    
    # Prepare context
    context_chunks = [result[0] for result in results]
    source_files = list(set([result[1] for result in results]))
    
    # Build prompt
    context = "\n\n".join(context_chunks) if context_chunks else "No relevant documents found."
    
    # Use agent's system prompt if available
    system_prompt = agent.system_prompt if agent else "You are a helpful legal assistant."
    
    # Include chat history
    history_context = ""
    if query_request.chat_history:
        history_context = "\n".join([
            f"{msg['role']}: {msg['content']}" 
            for msg in query_request.chat_history[-5:]  # Last 5 messages
        ])
    
    full_prompt = f"""
{system_prompt}

Context from documents:
{context}

Chat History:
{history_context}

Question: {query_request.question}

Please provide a helpful and accurate response based on the context provided.
"""
    
    # Generate response
    model_to_use = agent.model if agent else OLLAMA_MODEL
    use_gemini = "gemini" in model_to_use.lower()
    
    answer = get_llm_response(full_prompt, model_to_use, use_gemini)
    
    # Log messages
    user_message = ChatMessage(
        session_id=session_id,
        role="user",
        message=query_request.question,
        agent_name=agent_name,
        filtered=filter_result.get("keyword") if filter_result and filter_result["filter_type"] == "flag" else None
    )
    
    assistant_message = ChatMessage(
        session_id=session_id,
        role="assistant",
        message=answer,
        agent_name=agent_name
    )
    
    db.add(user_message)
    db.add(assistant_message)
    
    # Update session
    session.message_count += 2
    db.commit()
    
    return QueryResponse(
        answer=answer,
        sources=source_files,
        agent_used=agent_name,
        session_id=session_id
    )

# Include legal endpoints
from legal_endpoints import router as legal_router
app.include_router(legal_router)

# ============ STARTUP EVENT ============

@app.on_event("startup")
async def startup_event():
    # Create database tables
    Base.metadata.create_all(bind=engine)
    
    # Create default admin user if not exists
    db = SessionLocal()
    try:
        admin_user = db.query(User).filter(User.username == "admin").first()
        if not admin_user:
            admin_user = User(
                username="admin",
                email="admin@example.com",
                hashed_password=get_password_hash("admin123"),
                full_name="System Administrator",
                role=UserRole.SUPER_ADMIN,
                status=UserStatus.ACTIVE,
                is_verified=True
            )
            db.add(admin_user)
            db.commit()
            print("Created default admin user: admin/admin123")
        
        # Create default AI agents
        default_agents = [
            {
                "name": "general_legal_assistant",
                "alias": "General Legal Assistant",
                "model": OLLAMA_MODEL,
                "system_prompt": "You are a helpful legal assistant. Provide accurate legal information while reminding users that this is not legal advice and they should consult with a qualified attorney for specific legal matters.",
                "specialization": AgentSpecialization.GENERAL
            },
            {
                "name": "contract_analyzer",
                "alias": "Contract Analysis Specialist",
                "model": OLLAMA_MODEL,
                "system_prompt": "You are a contract analysis specialist. Help users understand contract terms, identify potential issues, and explain legal language in plain English.",
                "specialization": AgentSpecialization.CONTRACT_ANALYSIS
            }
        ]
        
        for agent_data in default_agents:
            existing_agent = db.query(AIAgent).filter(AIAgent.name == agent_data["name"]).first()
            if not existing_agent:
                agent = AIAgent(**agent_data)
                db.add(agent)
        
        db.commit()
        print("Database initialization completed")
        
    finally:
        db.close()

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)